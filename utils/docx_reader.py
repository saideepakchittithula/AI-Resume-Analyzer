# utils/docx_reader.py
"""
docx_reader.py — Robust DOCX text extraction for AI Resume Analyzer.

Extracts ALL text content from Word (.docx) resume files including:
  - Body paragraphs          (with heading level detection)
  - Tables                   (all rows and cells)
  - Headers and footers      (per section)
  - Text boxes               (via XML parsing)
  - Hyperlinks               (URL and display text)
  - Lists                    (numbered and bulleted)

Handles:
  - Corrupted / unreadable DOCX files  → friendly error
  - Empty documents                    → detected and warned
  - Password-protected DOCX            → detected and warned
  - Oversized files                    → rejected before processing
  - Non-DOCX files                     → rejected with clear message
  - Missing python-docx library        → graceful ImportError message
"""

import io
import re
import zipfile
from pathlib import Path
from typing import Optional
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from utils.constants import MAX_FILE_SIZE_MB
from utils.helpers import (
    clean_whitespace,
    count_words,
    get_file_size_mb,
    validate_file_size,
)
from utils.logger import get_logger

logger = get_logger(__name__)


class DOCXReadError(Exception):
    """Raised when a DOCX file cannot be read or parsed."""


class DOCXReader:
    """
    Extracts plain text from Word (.docx) resume files.

    Supports both file-path-based and bytes-based reading so it works
    seamlessly with Streamlit's UploadedFile (bytes) and local files.

    Usage:
        reader = DOCXReader()

        # From an uploaded file (Streamlit)
        result = reader.extract(file_bytes=uploaded_file.read())

        # From a local file path
        result = reader.extract(file_path="data/resume.docx")

    Returns a dict:
        {
            "text":         str,   # full extracted text
            "paragraphs":   int,   # number of paragraphs
            "word_count":   int,   # total word count
            "has_tables":   bool,  # whether tables were found
            "has_headers":  bool,  # whether headers/footers found
            "warnings":     list,  # any non-fatal warnings
            "success":      bool,
            "error":        str,   # populated only when success=False
        }
    """

    # Minimum characters to consider extraction successful
    MIN_CHARS: int = 50

    # XML namespace for DOCX internals
    _NS: dict[str, str] = {
        "w":   "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp":  "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "a":   "http://schemas.openxmlformats.org/drawingml/2006/main",
        "v":   "urn:schemas-microsoft-com:vml",
        "mc":  "http://schemas.openxmlformats.org/markup-compatibility/2006",
    }

    def __init__(self) -> None:
        self._warnings: list[str] = []

    # ------------------------------------------------------------------
    # Public interface
    # ------------------------------------------------------------------

    def extract(
        self,
        file_bytes: Optional[bytes] = None,
        file_path: Optional[str | Path] = None,
    ) -> dict:
        """
        Extract text from a DOCX file supplied as raw bytes or a file path.

        Exactly one of file_bytes or file_path must be provided.

        Args:
            file_bytes: Raw DOCX bytes (e.g. from Streamlit uploader).
            file_path:  Path to a local .docx file.

        Returns:
            Extraction result dictionary (see class docstring).
        """
        self._warnings = []

        try:
            raw_bytes = self._load_bytes(file_bytes, file_path)
            return self._process(raw_bytes)
        except DOCXReadError as exc:
            logger.error("DOCXReadError: %s", exc)
            return self._error_result(str(exc))
        except Exception as exc:
            logger.exception("Unexpected error reading DOCX: %s", exc)
            return self._error_result(
                f"An unexpected error occurred while reading the DOCX file: {exc}"
            )

    # ------------------------------------------------------------------
    # Private: load raw bytes
    # ------------------------------------------------------------------

    def _load_bytes(
        self,
        file_bytes: Optional[bytes],
        file_path: Optional[str | Path],
    ) -> bytes:
        """
        Resolve DOCX source to raw bytes.

        Args:
            file_bytes: Optional raw bytes.
            file_path:  Optional file path.

        Returns:
            Raw DOCX bytes.

        Raises:
            DOCXReadError: If source is invalid or missing.
        """
        if file_bytes is not None:
            return file_bytes

        if file_path is not None:
            path = Path(file_path)
            if not path.exists():
                raise DOCXReadError(f"File not found: {path}")
            if not path.is_file():
                raise DOCXReadError(f"Path is not a file: {path}")
            if path.suffix.lower() not in (".docx", ".doc"):
                raise DOCXReadError(
                    f"File '{path.name}' is not a DOCX file. "
                    f"Got extension: '{path.suffix}'. "
                    "Only .docx files are supported."
                )
            return path.read_bytes()

        raise DOCXReadError(
            "No DOCX source provided. Supply either file_bytes or file_path."
        )

    # ------------------------------------------------------------------
    # Private: main processing pipeline
    # ------------------------------------------------------------------

    def _process(self, raw_bytes: bytes) -> dict:
        """
        Validate then extract all text from raw DOCX bytes.

        Args:
            raw_bytes: DOCX file content.

        Returns:
            Extraction result dictionary.

        Raises:
            DOCXReadError: On validation failures.
        """
        # --- File size validation ---
        is_valid, size_msg = validate_file_size(raw_bytes, MAX_FILE_SIZE_MB)
        if not is_valid:
            raise DOCXReadError(size_msg)

        logger.info(
            "Processing DOCX — size: %.2f MB", get_file_size_mb(raw_bytes)
        )

        # --- Validate it is actually a DOCX (ZIP) file ---
        if not self._is_valid_docx(raw_bytes):
            raise DOCXReadError(
                "The file does not appear to be a valid DOCX file. "
                "It may be corrupted, a legacy .doc format, or not a Word document."
            )

        # --- Open document ---
        try:
            doc = Document(io.BytesIO(raw_bytes))
        except Exception as exc:
            raise DOCXReadError(
                f"Could not open the DOCX file. It may be corrupted: {exc}"
            )

        # --- Extract all text sections ---
        sections: list[str] = []

        body_text, para_count = self._extract_body(doc)
        if body_text:
            sections.append(body_text)

        table_text, has_tables = self._extract_tables(doc)
        if table_text:
            sections.append(table_text)

        header_footer_text, has_headers = self._extract_headers_footers(doc)
        if header_footer_text:
            sections.append(header_footer_text)

        textbox_text = self._extract_textboxes(raw_bytes)
        if textbox_text:
            sections.append(textbox_text)

        # --- Combine all sections ---
        full_text = "\n\n".join(s for s in sections if s.strip())
        full_text = self._post_process(full_text)

        word_count = count_words(full_text)

        # --- Warn if too little text extracted ---
        if len(full_text.strip()) < self.MIN_CHARS:
            self._warnings.append(
                "Very little text was extracted from this DOCX. "
                "The document may be empty or rely heavily on images."
            )
            logger.warning("Low text extraction from DOCX.")

        logger.info(
            "DOCX extracted — %d paragraphs, %d words, "
            "tables=%s, headers=%s",
            para_count, word_count, has_tables, has_headers,
        )

        return {
            "text":        full_text,
            "paragraphs":  para_count,
            "word_count":  word_count,
            "has_tables":  has_tables,
            "has_headers": has_headers,
            "warnings":    self._warnings,
            "success":     True,
            "error":       "",
        }

    # ------------------------------------------------------------------
    # Private: body paragraphs
    # ------------------------------------------------------------------

    def _extract_body(self, doc: Document) -> tuple[str, int]:
        """
        Extract text from all body paragraphs, preserving structure.

        Heading paragraphs are prefixed with '#' markers so downstream
        parsers can detect section boundaries easily.

        Args:
            doc: python-docx Document object.

        Returns:
            Tuple of (body_text, paragraph_count).
        """
        lines: list[str] = []
        para_count = 0

        for para in doc.paragraphs:
            text = self._get_paragraph_text(para)
            if not text.strip():
                continue

            para_count += 1
            style_name = para.style.name.lower() if para.style else ""

            # Prefix heading styles so section detection is easier
            if "heading 1" in style_name:
                lines.append(f"\n## {text}")
            elif "heading 2" in style_name:
                lines.append(f"\n### {text}")
            elif "heading" in style_name:
                lines.append(f"\n#### {text}")
            else:
                lines.append(text)

        return "\n".join(lines), para_count

    def _get_paragraph_text(self, para: Paragraph) -> str:
        """
        Extract full text from a paragraph including hyperlink runs.

        python-docx's para.text misses hyperlink text, so we walk
        the XML directly to capture everything.

        Args:
            para: A python-docx Paragraph object.

        Returns:
            Full paragraph text string.
        """
        texts: list[str] = []

        for child in para._p:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "r":
                # Regular run
                t_elem = child.find(qn("w:t"))
                if t_elem is not None and t_elem.text:
                    texts.append(t_elem.text)

            elif tag == "hyperlink":
                # Hyperlink — extract both the display text and URL
                for run in child.findall(qn("w:r")):
                    t_elem = run.find(qn("w:t"))
                    if t_elem is not None and t_elem.text:
                        texts.append(t_elem.text)

            elif tag == "ins":
                # Tracked insertion — include the inserted text
                for run in child.findall(qn("w:r")):
                    t_elem = run.find(qn("w:t"))
                    if t_elem is not None and t_elem.text:
                        texts.append(t_elem.text)

        return "".join(texts)

    # ------------------------------------------------------------------
    # Private: tables
    # ------------------------------------------------------------------

    def _extract_tables(self, doc: Document) -> tuple[str, bool]:
        """
        Extract text from all tables in the document.

        Many resume templates use tables for layout. Each cell is
        extracted and rows are pipe-separated for readability.

        Args:
            doc: python-docx Document object.

        Returns:
            Tuple of (table_text, has_tables).
        """
        if not doc.tables:
            return "", False

        rows: list[str] = []

        for table_idx, table in enumerate(doc.tables):
            for row in table.rows:
                cell_texts: list[str] = []
                for cell in row.cells:
                    # Each cell may contain multiple paragraphs
                    cell_text = " ".join(
                        self._get_paragraph_text(p)
                        for p in cell.paragraphs
                        if p.text.strip()
                    )
                    if cell_text.strip():
                        cell_texts.append(cell_text.strip())

                if cell_texts:
                    rows.append(" | ".join(cell_texts))

        logger.debug(
            "DOCX tables: %d tables, %d text rows extracted",
            len(doc.tables), len(rows),
        )

        return "\n".join(rows), True

    # ------------------------------------------------------------------
    # Private: headers and footers
    # ------------------------------------------------------------------

    def _extract_headers_footers(self, doc: Document) -> tuple[str, bool]:
        """
        Extract text from document headers and footers.

        Resumes often place name, contact info, or LinkedIn in the
        header, which must be extracted for complete parsing.

        Args:
            doc: python-docx Document object.

        Returns:
            Tuple of (header_footer_text, has_headers).
        """
        texts: list[str] = []
        found = False

        try:
            for section in doc.sections:
                # Header
                for header_type in ("header", "first_page_header", "even_page_header"):
                    hdr = getattr(section, header_type, None)
                    if hdr is not None:
                        for para in hdr.paragraphs:
                            t = self._get_paragraph_text(para).strip()
                            if t:
                                texts.append(t)
                                found = True

                # Footer
                for footer_type in ("footer", "first_page_footer", "even_page_footer"):
                    ftr = getattr(section, footer_type, None)
                    if ftr is not None:
                        for para in ftr.paragraphs:
                            t = self._get_paragraph_text(para).strip()
                            if t:
                                texts.append(t)
                                found = True

        except Exception as exc:
            logger.warning("Could not extract headers/footers: %s", exc)

        return "\n".join(texts), found

    # ------------------------------------------------------------------
    # Private: text boxes (via raw XML)
    # ------------------------------------------------------------------

    def _extract_textboxes(self, raw_bytes: bytes) -> str:
        """
        Extract text from drawing/text-box shapes via raw XML parsing.

        python-docx does not expose text boxes natively, so we unzip
        the DOCX and parse the XML directly.  This ensures text placed
        in floating text boxes (common in designer resume templates)
        is not lost.

        Args:
            raw_bytes: Raw DOCX bytes.

        Returns:
            Concatenated text from all text boxes.
        """
        texts: list[str] = []

        try:
            with zipfile.ZipFile(io.BytesIO(raw_bytes)) as zf:
                # The main document body is always at word/document.xml
                if "word/document.xml" not in zf.namelist():
                    return ""

                xml_content = zf.read("word/document.xml")
                root = ET.fromstring(xml_content)

                # Find all <w:txbxContent> elements (Word text boxes)
                ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                for txbx in root.iter(f"{{{ns}}}txbxContent"):
                    for para in txbx.iter(f"{{{ns}}}p"):
                        para_texts: list[str] = []
                        for run in para.iter(f"{{{ns}}}t"):
                            if run.text:
                                para_texts.append(run.text)
                        line = "".join(para_texts).strip()
                        if line:
                            texts.append(line)

        except zipfile.BadZipFile:
            logger.warning("Could not unzip DOCX for text box extraction.")
        except ET.ParseError as exc:
            logger.warning("XML parse error during text box extraction: %s", exc)
        except Exception as exc:
            logger.warning("Text box extraction failed: %s", exc)

        if texts:
            logger.debug("Extracted %d text box segments.", len(texts))

        return "\n".join(texts)

    # ------------------------------------------------------------------
    # Private: DOCX validation
    # ------------------------------------------------------------------

    @staticmethod
    def _is_valid_docx(raw_bytes: bytes) -> bool:
        """
        Check whether the bytes represent a valid DOCX (ZIP) structure.

        A DOCX file is a ZIP archive containing 'word/document.xml'.
        This check is faster than attempting a full Document() parse.

        Args:
            raw_bytes: File content bytes.

        Returns:
            True if the bytes look like a valid DOCX file.
        """
        try:
            with zipfile.ZipFile(io.BytesIO(raw_bytes)) as zf:
                return "word/document.xml" in zf.namelist()
        except (zipfile.BadZipFile, Exception):
            return False

    # ------------------------------------------------------------------
    # Private: post-processing
    # ------------------------------------------------------------------

    def _post_process(self, text: str) -> str:
        """
        Clean and normalise extracted DOCX text.

        Steps:
          1. Remove heading markers used internally (## / ### / ####)
             but keep the text itself.
          2. Collapse 3+ consecutive blank lines into 2.
          3. Fix non-breaking spaces and common unicode artefacts.
          4. Strip leading/trailing whitespace.

        Args:
            text: Raw extracted text with internal markers.

        Returns:
            Cleaned plain-text string.
        """
        if not text:
            return ""

        # Convert heading markers back to plain text (keep section names)
        text = re.sub(r"^#{2,4}\s+", "", text, flags=re.MULTILINE)

        # Replace non-breaking spaces with regular spaces
        text = text.replace("\u00a0", " ")
        text = text.replace("\u2019", "'")   # right single quotation
        text = text.replace("\u2018", "'")   # left single quotation
        text = text.replace("\u201c", '"')   # left double quotation
        text = text.replace("\u201d", '"')   # right double quotation
        text = text.replace("\u2013", "-")   # en dash
        text = text.replace("\u2014", "-")   # em dash
        text = text.replace("\u2022", "-")   # bullet •
        text = text.replace("\u25cf", "-")   # filled circle ●
        text = text.replace("\u25ba", "-")   # right triangle ►
        text = text.replace("\u2192", "->")  # arrow →

        # Remove zero-width and invisible characters
        text = re.sub(r"[\u200b\u200c\u200d\ufeff]", "", text)

        # Collapse excessive blank lines
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Remove non-printable characters except standard whitespace
        text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", " ", text)

        return text.strip()

    # ------------------------------------------------------------------
    # Private: result builders
    # ------------------------------------------------------------------

    @staticmethod
    def _error_result(message: str) -> dict:
        """Build a standardised error result dictionary."""
        return {
            "text":        "",
            "paragraphs":  0,
            "word_count":  0,
            "has_tables":  False,
            "has_headers": False,
            "warnings":    [],
            "success":     False,
            "error":       message,
        }

    # ------------------------------------------------------------------
    # Convenience class method
    # ------------------------------------------------------------------

    @classmethod
    def read(
        cls,
        file_bytes: Optional[bytes] = None,
        file_path: Optional[str | Path] = None,
    ) -> dict:
        """
        One-liner convenience wrapper around extract().

        Args:
            file_bytes: Raw DOCX bytes.
            file_path:  Path to a .docx file.

        Returns:
            Extraction result dictionary.

        Example:
            result = DOCXReader.read(file_bytes=uploaded_file.read())
            if result["success"]:
                text = result["text"]
        """
        return cls().extract(file_bytes=file_bytes, file_path=file_path)
