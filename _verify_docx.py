import sys, io, importlib.util
sys.path.insert(0, ".")

from docx import Document
from docx.shared import Pt

def load(name, path):
    spec = importlib.util.spec_from_file_location(name, path)
    m = importlib.util.module_from_spec(spec)
    sys.modules[name] = m
    spec.loader.exec_module(m)
    return m

load("utils.constants", "utils/constants.py")
load("utils.logger",    "utils/logger.py")
load("utils.helpers",   "utils/helpers.py")
docx_mod = load("utils.docx_reader", "utils/docx_reader.py")

DOCXReader    = docx_mod.DOCXReader
DOCXReadError = docx_mod.DOCXReadError

reader = DOCXReader()
print("DOCXReader instantiated successfully")
print()

# -------------------------------------------------------
# Helper: build a real DOCX in memory with rich content
# -------------------------------------------------------
def make_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("John Doe", level=1)
    doc.add_paragraph("john.doe@email.com | +1 555-867-5309 | linkedin.com/in/johndoe | github.com/johndoe")
    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(
        "Senior Python Developer with 6 years of experience building "
        "scalable AI systems using FastAPI, LangChain, and AWS."
    )
    doc.add_heading("Skills", level=2)
    doc.add_paragraph("Python, FastAPI, Docker, AWS, LangChain, PostgreSQL, React")
    doc.add_heading("Experience", level=2)
    doc.add_paragraph("Senior AI Engineer — TechCorp (2021 - Present)")
    doc.add_paragraph("- Built LLM-powered pipelines reducing processing time by 40%")
    doc.add_paragraph("- Deployed microservices on AWS ECS with CI/CD via GitHub Actions")
    doc.add_heading("Education", level=2)
    doc.add_paragraph("B.Sc. Computer Science — MIT, 2018")
    doc.add_heading("Projects", level=2)
    doc.add_paragraph("AI Resume Analyzer — Python, Streamlit, spaCy | github.com/johndoe/ai-resume")

    # Add a table (common in resume templates)
    table = doc.add_table(rows=2, cols=3)
    table.cell(0,0).text = "Python"
    table.cell(0,1).text = "FastAPI"
    table.cell(0,2).text = "Docker"
    table.cell(1,0).text = "AWS"
    table.cell(1,1).text = "LangChain"
    table.cell(1,2).text = "PostgreSQL"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

docx_bytes = make_docx_bytes()
print(f"Sample DOCX built in memory: {len(docx_bytes)} bytes")
print()

# -------------------------------------------------------
# Test 1: Successful extraction
# -------------------------------------------------------
result = reader.extract(file_bytes=docx_bytes)
assert result["success"] == True, f"Expected success, got: {result['error']}"
assert len(result["text"]) > 100
assert result["word_count"] > 20
assert result["paragraphs"] > 0
assert result["has_tables"] == True
print("Test 1 PASSED: Full DOCX extraction")
print(f"         paragraphs={result['paragraphs']}, words={result['word_count']}, tables={result['has_tables']}")
print(f"         Text preview: {result['text'][:80].strip()}...")
print()

# -------------------------------------------------------
# Test 2: Key content present in extracted text
# -------------------------------------------------------
text = result["text"]
assert "John Doe"             in text, "Name not found"
assert "john.doe@email.com"   in text, "Email not found"
assert "Python"               in text, "Skill not found"
assert "FastAPI"              in text, "Skill not found"
assert "TechCorp"             in text, "Company not found"
assert "MIT"                  in text, "Education not found"
print("Test 2 PASSED: All key content present in extracted text")
print()

# -------------------------------------------------------
# Test 3: No source provided
# -------------------------------------------------------
result = reader.extract()
assert result["success"] == False
assert "No DOCX source" in result["error"]
print("Test 3 PASSED: No source provided ->", result["error"][:55])

# -------------------------------------------------------
# Test 4: Non-existent file path
# -------------------------------------------------------
result = reader.extract(file_path="data/nonexistent.docx")
assert result["success"] == False
assert "not found" in result["error"].lower()
print("Test 4 PASSED: Nonexistent file  ->", result["error"][:55])

# -------------------------------------------------------
# Test 5: Wrong file extension
# -------------------------------------------------------
result = reader.extract(file_path="data/sample_jd.txt")
assert result["success"] == False
assert "not a DOCX" in result["error"]
print("Test 5 PASSED: Wrong extension   ->", result["error"][:55])

# -------------------------------------------------------
# Test 6: Oversized file
# -------------------------------------------------------
big = b"PK" + b"x" * (11 * 1024 * 1024)
result = reader.extract(file_bytes=big)
assert result["success"] == False
assert "exceeds" in result["error"]
print("Test 6 PASSED: File too large    ->", result["error"][:55])

# -------------------------------------------------------
# Test 7: Corrupted bytes
# -------------------------------------------------------
result = reader.extract(file_bytes=b"this is not a docx file")
assert result["success"] == False
print("Test 7 PASSED: Corrupted bytes   ->", result["error"][:55])

# -------------------------------------------------------
# Test 8: is_valid_docx check
# -------------------------------------------------------
assert DOCXReader._is_valid_docx(docx_bytes) == True
assert DOCXReader._is_valid_docx(b"not a zip") == False
print("Test 8 PASSED: _is_valid_docx works correctly")

# -------------------------------------------------------
# Test 9: Post-process cleanup
# -------------------------------------------------------
raw = "## Summary\nHello\u00a0World\u2013Test\n\n\n\nExtra"
cleaned = reader._post_process(raw)
assert "##" not in cleaned
assert "\u00a0" not in cleaned
assert "\u2013" not in cleaned
print("Test 9 PASSED: Post-process      ->", repr(cleaned[:40]))

# -------------------------------------------------------
# Test 10: Class method shortcut
# -------------------------------------------------------
result2 = DOCXReader.read(file_bytes=docx_bytes)
assert result2["success"] == True
assert result2["text"] == result["text"]
print("Test 10 PASSED: DOCXReader.read() class method works")

# -------------------------------------------------------
# Test 11: Error result structure
# -------------------------------------------------------
err = DOCXReader._error_result("test error")
assert set(err.keys()) == {
    "text","paragraphs","word_count","has_tables",
    "has_headers","warnings","success","error"
}
print("Test 11 PASSED: Error result dict has all keys")

print()
print("ALL DOCX_READER.PY TESTS PASSED")
